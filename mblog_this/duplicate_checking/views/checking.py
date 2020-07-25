"""
author:syz
time:2020-4-13
"""


"""
查重步骤：

1.读取word文档，按照单元格来读
2.获取文章内容列表（每个单元格，一个位置，去掉重复的），按照每行获取文章内容字典，将列表存入redis，以便后期分词。（可以
按照不同的行获取相应行的内容，后期加上）
3.按照字典，正则保留汉子，求出汉子总个数。
4.按照列表进行分词。将分词后的列表加进总分词列表。
5.对总分词列表产出词袋
6.对总分词表转化为id表示的文档向量

"""
import os

import docx

import re
from collections import Counter
from duplicate_checking.views.decorators import aim_or_whole
import jieba
import warnings
warnings.filterwarnings(action='ignore', category=UserWarning, module='gensim')
from gensim import corpora,models,similarities



# from django_redis import get_redis_connection

# redis = get_redis_connection('default')




class WordOperationBase:
    # 例如 '软件工程-1705-司云中-17408002125-Python课程-实验2'
    # filename 格式 专业-班级-姓名-学号-课程

    filename_regex = r'([\u4e00-\u9fa5]{2,10})-(\d{4})-(\d{11})-([\u4e00-\u9fa5]{2,5})-' \
                     r'([\u4e00-\u9fa5]{2,20}\w*|\w*[\u4e00-\u9fa5]{2,20})-(实验\d{1,2})'
    _name = None
    _class = None
    _major = None
    _number = None
    _project = None
    _experiment = None
    _table = None
    _doc = None  # docx实例
    _disqualification = None  # 失败的人员

    def __init__(self, filename,rows=None,paragraph=None):
        self._filename = os.path.join(r'D:\\syz\\virtualenvs\\MyDjango\\svn_mblog\\mblog\\', filename)
        self.split_filename() # 分割文件名
        setattr(self,self.doc,self.get_doc())
        setattr(self, self.table, self.get_tables())

    @property
    def name(self):
        return getattr(self,self._name)

    @name.setter
    def name(self,value):
        setattr(self,self._name,value)

    @property
    def major(self):
        return getattr(self,self._major)

    @major.setter
    def major(self,value):
        setattr(self,self._major,value)

    @property
    def project(self):
        return getattr(self,self._project)

    @project.setter
    def project(self, value):
        setattr(self,self._project,value)

    @property
    def number(self):
        return getattr(self,self._number)

    @number.setter
    def number(self, value):
        setattr(self,self._number,value)

    @property
    def experiment(self):
        return getattr(self,self._experiment)

    @experiment.setter
    def experiment(self,value):
        setattr(self,self._experiment,value)

    @property
    def class_(self):
        return getattr(self,self._class)

    @class_.setter
    def class_(self, value):
        setattr(self,self._class,value)

    @property
    def table(self):
        return getattr(self,self._table)

    @table.setter
    def table(self, value):
        setattr(self,self._table,value)

    @property
    def doc(self):
        return getattr(self,self._doc)

    @doc.setter
    def doc(self, value):
        setattr(self,self._doc,value)

    @property
    def disqualification(self):
        return self._disqualification

    @disqualification.setter
    def disqualification(self,value):
        self._disqualification = value


    def __repr__(self):
        return '%s-%s' % (self.number,self.name)

    def create_pattern(self):
        pattern = re.compile(self.filename_regex, re.I)
        return pattern

    def get_doc(self):
        """获取doc对象"""
        return docx.Document(self._filename)
        # fullText = [para.text for para in self.doc.paragraphs]
        # return '\n'.join(fullText)

    def split_filename(self):
        """分割匹配集"""
        pattern = self.create_pattern()
        result = pattern.findall(self._filename)[0]
        self.major = result[0]
        self.class_ = result[1]
        self.number = result[2]
        self.name = result[3]
        self.project = result[4]
        self.experment = result[5]

    def get_tables(self, index=None):
        """
        默认第一个表格
        index :表格的索引
        """
        if not index:
            return self.doc.tables[0]
        elif isinstance(index, int):
            pass  # 异常检测
        else:
            return self.doc.tables[index]

    def get_tb_rows(self):
        """获取表格行数"""
        return len(self.table.rows)

    def get_one_row(self, row_index):
        """获取某一行"""
        return self.table.rows[row_index]

    def get_tb_cells(self, row_index):
        """获取某一行的列数"""
        return len(self.table.rows[row_index].cells)

    def get_aim_content(self, row=None, **kwargs):
        """获取目标内容"""
        global rows
        dict_context = {}
        list_context = []
        if not row:
            rows = self.table.rows
            for index, value in enumerate(rows):
                dict_context.setdefault(index, [])
            for index_r, value_r in enumerate(rows):
                for index_c, value_c in enumerate(value_r.cells):
                    # text = re.sub(self.content_regex,value_r.text,'')
                    text = value_c.text.replace('\n', '')
                    if text not in list_context:
                        dict_context[index_r].append(text)
                        list_context.append(text)
                    else:
                        # 去掉重复单元格数据，减少一定的时间复杂度
                        break
        else:
            rows = self.get_one_row(row)
            pass
            # 自定义
        # 可以存入redis中
        self.list_context = list_context
        self.dict_context = dict_context
        # redis.set(self.get_word_key(),list_context)

    @aim_or_whole(True)
    def words_counts(self, rows=None ,whole=None):
        """
        计算汉字个数
        row:行索引列表
        cells:列索引
        子类可以重写,个性化定制row和cell进行计算
        """
        if not rows:
            # 匹配所有行
            self.get_aim_content(rows)
            return self.get_dict_word()
        else:
            # 匹配某些行表格行，自定义
            pass

    def compute_words_counts(self):
        self.words_counts = self.words_counts()

    def get_words_counts(self):
        """获取总字数"""
        return getattr(self,self.words_counts)

    def get_disqualification(self):
        """获取不合格的同学列表 """
        return self.disqualification

    def get_word_key(self):
        """返回文章唯一标识"""
        self.word_key = '%s-%s' % (self.number, self.name)
        return self.word_key

    def get_list_word(self):
        """获取文章每单元格内容列表"""
        return self.list_context

    def get_dict_word(self):
        """获取文章每行内容字典"""
        return self.dict_context


class Participle(WordOperationBase):
    """分词，提取关键词"""
    def __init__(self,filename):
        super().__init__(filename)
        self.compute_words_counts()  # 获取

    def clean_get_word(self):
        """对本篇实验报告进行分词和清洗字符串"""
        # redis读取相应用户的关键词
        # context = redis.get(word_key)
        context = self.get_list_word()
        self.context_str = ''.join(context)
        pattern  = r'[^\u4e00-\u9fff]+'
        self.context_str = re.sub(pattern,'',self.context_str)
        self.particle_words = jieba.cut(self.context_str,True)

    def add_particle_word(self,total_particle_list):
        """添加分词列表到总的分词列表中"""
        self.clean_get_word()
        particle_words = [word for word in self.particle_words] # 获取分词列表
        total_particle_list.append(particle_words) # 将分词表添加到总的分词列表

    def get_context_str(self):
        """获取内容字符串格式"""
        return self.context_str

    def get_particle_words(self):
        """获取分词包,生成器类型"""
        return self.particle_words


class Duplicate_checking():

    def get_bag_words(self):
        """获取词袋，token作为key，词组作为value"""
        dictionary = corpora.Dictionary(m)
        self.bag_words = dictionary.token2id
        pass

    def trans_vector(self):
        """将字符串文档转为id表示的文档向量"""
        pass



file = Participle('软件工程-1705-17408002125-司云中-JavaEE开发技术-实验2.docx')



